"""
API endpoints для шаблонов документов и генерации
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from typing import Optional
from datetime import datetime
from pathlib import Path
import aiofiles
from app.database import get_db
from app.models.user import User
from app.models.document_template import DocumentTemplate
from app.models.case import Case
from app.schemas.document_template import (
    DocumentTemplateResponse, DocumentTemplateListResponse
)
from app.api.deps import get_current_user
from app.config import settings
from app.utils.ollama import ollama_client
import logging

logger = logging.getLogger(__name__)

router = APIRouter()


@router.get("/", response_model=DocumentTemplateListResponse)
async def get_templates(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Получение списка всех шаблонов"""
    result = await db.execute(
        select(DocumentTemplate).order_by(DocumentTemplate.created_at.desc())
    )
    templates = result.scalars().all()

    return DocumentTemplateListResponse(
        items=templates,
        total=len(templates)
    )


@router.post("/", response_model=DocumentTemplateResponse, status_code=status.HTTP_201_CREATED)
async def upload_template(
    file: UploadFile = File(...),
    template_name: str = Form(...),
    template_type: str = Form(...),
    description: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Загрузка нового шаблона документа (DOCX)"""
    # Проверка формата
    if not file.filename.endswith('.docx'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Поддерживаются только DOCX файлы"
        )

    # Сохранение файла
    storage_path = Path(settings.STORAGE_PATH) / "templates"
    storage_path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    new_filename = f"{timestamp}_{file.filename.replace(' ', '_')}"
    file_path = storage_path / new_filename

    content = await file.read()
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)

    # Извлечение переменных из шаблона ({{variable_name}})
    import re
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        variables = list(set(re.findall(r'\{\{(\w+)\}\}', text)))
    except Exception as e:
        logger.warning(f"Не удалось извлечь переменные: {e}")
        variables = []

    # Создание записи
    new_template = DocumentTemplate(
        template_name=template_name,
        template_type=template_type,
        description=description,
        file_path=f"templates/{new_filename}",
        file_size=len(content),
        variables={"variables": variables} if variables else None,
        created_by=current_user.id
    )

    db.add(new_template)
    await db.commit()
    await db.refresh(new_template)

    return new_template


@router.post("/{template_id}/generate")
async def generate_document(
    template_id: int,
    case_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Генерация документа по шаблону с заполнением данных дела
    """
    # Получение шаблона
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()

    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Шаблон не найден")

    # Получение дела
    result = await db.execute(select(Case).where(Case.id == case_id))
    case = result.scalar_one_or_none()

    if not case:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Дело не найдено")

    # Генерация документа
    from docx import Document as DocxDocument

    template_path = Path(settings.STORAGE_PATH) / template.file_path
    doc = DocxDocument(template_path)

    # Замена переменных
    replacements = {
        "case_number": case.case_number,
        "case_title": case.title,
        "plaintiff": case.plaintiff or "",
        "defendant": case.defendant or "",
        "court": case.court or "",
        "judge": case.judge or "",
        "open_date": case.open_date.strftime('%d.%m.%Y') if case.open_date else ""
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

    # Сохранение сгенерированного документа
    output_dir = Path(settings.STORAGE_PATH) / "documents" / f"case_{case_id}"
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f"{timestamp}_generated_{template.template_name}.docx"
    output_path = output_dir / output_filename

    doc.save(output_path)

    return {
        "success": True,
        "file_path": f"case_{case_id}/{output_filename}",
        "message": "Документ успешно сгенерирован"
    }


@router.delete("/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_template(
    template_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Удаление шаблона (только admin)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Требуются права администратора")

    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()

    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Шаблон не найден")

    # Удаление файла
    try:
        file_path = Path(settings.STORAGE_PATH) / template.file_path
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        logger.error(f"Ошибка при удалении файла: {e}")

    await db.execute(delete(DocumentTemplate).where(DocumentTemplate.id == template_id))
    await db.commit()

    return None
